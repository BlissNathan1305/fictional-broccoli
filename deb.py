import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# === Data ===
banana_sap_composition = {
    "Moisture (%)": 95.81,
    "Protein (%)": 1.75,
    "Fat/Lipid (%)": 0.23,
    "Fibre (%)": 0.00,
    "Ash (%)": 0.12,
    "Carbohydrate (%)": 2.08,
    "Energy (kcal/100g)": 17.39,
    "Lignin (%)": 0.01,
    "Hemicellulose (%)": 0.51,
    "Cellulose (%)": 0.58,
    "Sugar (%)": 5.13
}

# === Derived metrics ===
moisture = banana_sap_composition["Moisture (%)"]
dry_matter = 100 - moisture
sugar = banana_sap_composition["Sugar (%)"]
energy = banana_sap_composition["Energy (kcal/100g)"]

energy_per_sugar = energy / sugar if sugar != 0 else 0
ethanol_yield = sugar * 0.51  # 0.51 g ethanol per g sugar

metrics = {
    "Moisture (%)": moisture,
    "Dry matter (%)": dry_matter,
    "Sugar (g / 100 g fresh)": sugar,
    "Energy (kcal / 100 g fresh)": energy,
    "Energy per sugar (kcal per g sugar)": energy_per_sugar,
    "Theoretical ethanol yield (g/100 g fresh)": ethanol_yield,
}

# === Output directory ===
output_dir = "banana_sap_analysis"
os.makedirs(output_dir, exist_ok=True)

# === Plots ===
# Proximate composition: moisture → carbohydrate
proximate = {k: v for k, v in banana_sap_composition.items() if k in [
    "Moisture (%)", "Protein (%)", "Fat/Lipid (%)", "Fibre (%)", "Ash (%)", "Carbohydrate (%)"
]}
plt.figure(figsize=(8, 5))
plt.bar(proximate.keys(), proximate.values(), color="skyblue")
plt.title("Proximate Composition of Banana Sap")
plt.ylabel("Percentage (%)")
plt.xticks(rotation=45, ha="right")
plt.tight_layout()
proximate_path = os.path.join(output_dir, "proximate_composition.jpg")
plt.savefig(proximate_path, dpi=300)
plt.close()

# Bioethanol relevant metrics: energy → sugar
bioethanol = {k: v for k, v in banana_sap_composition.items() if k in [
    "Energy (kcal/100g)", "Lignin (%)", "Hemicellulose (%)", "Cellulose (%)", "Sugar (%)"
]}
plt.figure(figsize=(8, 5))
plt.bar(bioethanol.keys(), bioethanol.values(), color="orange")
plt.title("Bioethanol-Relevant Metrics of Banana Sap")
plt.ylabel("Value")
plt.xticks(rotation=45, ha="right")
plt.tight_layout()
bioethanol_path = os.path.join(output_dir, "bioethanol_metrics.jpg")
plt.savefig(bioethanol_path, dpi=300)
plt.close()

# === Word report ===
doc = Document()
doc.add_heading("Statistical and Bioethanol Analysis of Banana Sap", 0)

# Overall composition table
doc.add_heading("General Composition Data", level=1)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Metric"
hdr_cells[1].text = "Value"
for k, v in banana_sap_composition.items():
    row_cells = table.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = str(v)

# Derived metrics table
doc.add_heading("Derived Metrics", level=1)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Metric"
hdr_cells[1].text = "Value"
for k, v in metrics.items():
    row_cells = table.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = f"{v:.2f}"

# Proximate composition section
doc.add_heading("Proximate Composition", level=1)

# Table for proximate composition
doc.add_paragraph("The proximate composition values are presented below:")
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Component"
hdr_cells[1].text = "Percentage (%)"
for k, v in proximate.items():
    row_cells = table.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = f"{v:.2f}"

# Graph
doc.add_picture(proximate_path, width=Inches(5))

# Discussion
doc.add_heading("Discussion: Proximate Composition", level=2)
doc.add_paragraph(
    f"The proximate composition shows that banana sap is overwhelmingly composed "
    f"of moisture ({moisture:.2f}%%), leaving only {dry_matter:.2f}%% dry matter. "
    f"Within the dry matter, protein ({banana_sap_composition['Protein (%)']:.2f}%%), "
    f"lipid ({banana_sap_composition['Fat/Lipid (%)']:.2f}%%), ash "
    f"({banana_sap_composition['Ash (%)']:.2f}%%), and fibre are present only in trace "
    f"amounts. Carbohydrates ({banana_sap_composition['Carbohydrate (%)']:.2f}%%) "
    f"constitute the main component of the solid fraction, which is important for "
    f"nutritional and biochemical relevance."
)

# Bioethanol metrics section
doc.add_heading("Bioethanol-Relevant Metrics", level=1)

# Table for bioethanol metrics
doc.add_paragraph("The bioethanol-relevant metrics are summarized below:")
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Metric"
hdr_cells[1].text = "Value"
for k, v in bioethanol.items():
    row_cells = table.add_row().cells
    row_cells[0].text = k
    row_cells[1].text = f"{v:.2f}"

# Graph
doc.add_picture(bioethanol_path, width=Inches(5))

# Discussion
doc.add_heading("Discussion: Bioethanol-Relevant Metrics", level=2)
doc.add_paragraph(
    f"Among the bioethanol-relevant metrics, sugars ({sugar:.2f}%%) are the most "
    f"significant fraction, directly contributing to fermentable substrate. The "
    f"energy content ({energy:.2f} kcal/100 g) is modest but aligns with the high "
    f"water fraction. Lignin ({banana_sap_composition['Lignin (%)']:.2f}%%), "
    f"hemicellulose ({banana_sap_composition['Hemicellulose (%)']:.2f}%%), and "
    f"cellulose ({banana_sap_composition['Cellulose (%)']:.2f}%%) are present in "
    f"trace amounts, suggesting limited structural biomass but minimal inhibitory "
    f"effects for fermentation. The calculated ethanol yield of {ethanol_yield:.2f} g "
    f"per 100 g sap underscores its potential in bioethanol applications."
)

# Conclusion
doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    f"In conclusion, banana sap is predominantly water but contains a fermentable "
    f"sugar fraction that makes it relevant in bioethanol research. While its proximate "
    f"composition indicates low nutritional value beyond hydration, the bioethanol-relevant "
    f"metrics highlight a modest but usable sugar source. The low lignin and cellulose "
    f"contents also suggest that fermentation can proceed with minimal pretreatment, making "
    f"banana sap an interesting supplementary substrate for sustainable biofuel production."
)

# Save document
doc_path = os.path.join(output_dir, "banana_sap_analysis_report.docx")
doc.save(doc_path)

print("Analysis complete. Files saved in:", output_dir)
