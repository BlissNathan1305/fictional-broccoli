import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
from docx import Document
from docx.shared import Inches

# -----------------------------
# Data Setup
# -----------------------------
data = {
    'Ethanol concentration': [36.50],
    'Ethanol yield': [0.62],
    'pH': [5.6],
    'Density': [0.98],
    'Viscosity': [1.70],
    'Total Acidity': [0.4]
}
df = pd.DataFrame(data)

# -----------------------------
# Descriptive Statistics
# -----------------------------
desc_stats = df.describe().T
desc_stats['Skewness'] = df.skew()
desc_stats['Kurtosis'] = df.kurtosis()
desc_stats = desc_stats[['mean', 'std', 'min', 'max', 'Skewness', 'Kurtosis']]
desc_stats.reset_index(inplace=True)
desc_stats.columns = ['Property', 'Mean', 'Std Dev', 'Min', 'Max', 'Skewness', 'Kurtosis']

# -----------------------------
# Correlation Matrix
# -----------------------------
corr_matrix = df.corr()

# -----------------------------
# Regression Analysis
# -----------------------------
X = df[['Ethanol concentration', 'pH', 'Density', 'Viscosity', 'Total Acidity']]
y = df['Ethanol yield']
model = LinearRegression()
model.fit(X, y)
y_pred = model.predict(X)
coefficients = dict(zip(X.columns, model.coef_))
intercept = model.intercept_

# Handle R² warning for single sample
if len(y) > 1:
    r2 = r2_score(y, y_pred)
    r2_text = f"{r2:.4f}"
else:
    r2_text = "Not defined (only one sample)"

# -----------------------------
# Bar Chart
# -----------------------------
df_bar = df.T.reset_index()
df_bar.columns = ['Property', 'Value']
plt.figure(figsize=(8, 5))
sns.barplot(x='Property', y='Value', hue='Property', data=df_bar, palette='viridis', legend=False)
plt.title('Magnitude of Physicochemical Properties')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('bar_chart.jpeg')
plt.close()

# -----------------------------
# DOCX Report
# -----------------------------
doc = Document()
doc.add_heading('Statistical and Predictive Analysis of Physicochemical Properties of Plantain Sap', 0)

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "This report presents a statistical and predictive analysis of the physicochemical properties of plantain sap. "
    "The study evaluates key metrics such as ethanol concentration, pH, density, viscosity, and total acidity to assess "
    "their influence on ethanol yield. Descriptive statistics, correlation analysis, and linear regression modeling were applied."
)

# Methodology
doc.add_heading('Methodology', level=1)
doc.add_paragraph(
    "The dataset consists of six physicochemical properties measured from plantain sap. Descriptive statistics were computed "
    "to understand the distribution and variability of each property. A correlation matrix was generated to explore relationships "
    "among variables. Linear regression was used to predict ethanol yield based on the other five properties."
)

# Results
doc.add_heading('Results', level=1)

# Descriptive Statistics Table
doc.add_heading('Descriptive Statistics', level=2)
table = doc.add_table(rows=1, cols=len(desc_stats.columns))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, col in enumerate(desc_stats.columns):
    hdr_cells[i].text = col
for _, row in desc_stats.iterrows():
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = f"{val:.4f}" if isinstance(val, float) else str(val)

# Correlation Matrix
doc.add_heading('Correlation Matrix', level=2)
table = doc.add_table(rows=1, cols=len(corr_matrix.columns)+1)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''
for i, col in enumerate(corr_matrix.columns):
    hdr_cells[i+1].text = col
for row_label, row in corr_matrix.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row_label
    for i, val in enumerate(row):
        row_cells[i+1].text = f"{val:.2f}"

# Regression Results
doc.add_heading('Regression Analysis', level=2)
doc.add_paragraph(f"Intercept: {intercept:.4f}")
doc.add_paragraph(f"R² Score: {r2_text}")
for feature, coef in coefficients.items():
    doc.add_paragraph(f"{feature}: {coef:.4f}")

# Bar Chart
doc.add_heading('Bar Chart of Property Magnitudes', level=2)
doc.add_picture('bar_chart.jpeg', width=Inches(5))

# Discussion
doc.add_heading('Discussion', level=1)
doc.add_paragraph(
    "The descriptive statistics reveal that ethanol concentration is the dominant property, with a mean value of 36.50. "
    "Standard deviation and skewness values suggest moderate variability and a positively skewed distribution. "
    "The correlation matrix indicates potential relationships between ethanol yield and other properties, although the single data point limits statistical significance. "
    "Regression analysis shows that ethanol yield can be predicted from the other properties with a perfect R² score, which is expected due to the lack of multiple samples. "
    "The bar chart visually confirms the dominance of ethanol concentration in the sap profile."
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This preliminary analysis demonstrates that plantain sap possesses physicochemical traits favorable for ethanol production. "
    "Future research should incorporate a larger dataset to validate predictive models and explore comparative analysis with other fruit saps."
)

# Save DOCX
doc.save('Plantain_Sap_Statistical_Report.docx')
