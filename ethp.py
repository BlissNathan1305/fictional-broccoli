import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from math import pi

# Step 1: Define the data
data = {
    "Octane number": (90, 110),
    "Flash point": (13, 16),
    "Density": (0.79, 0.81),
    "Viscosity": (1.1, 1.4),
    "Vapor pressure": (5.5, 6.2),
    "Calorific value": (26.5, 27.5),
    "Boiling point": (78, 78),
    "Freezing point": (-114, -95),
    "Autoignition temp.": (360, 370)
}

# Step 2: Categorize properties
thermal = ["Flash point", "Boiling point", "Freezing point", "Autoignition temp."]
performance = ["Octane number", "Calorific value"]
physical = ["Density", "Viscosity", "Vapor pressure"]

# Step 3: Create plots
def plot_bar(data_subset, title, filename):
    labels = list(data_subset.keys())
    mins = [data_subset[k][0] for k in labels]
    maxs = [data_subset[k][1] for k in labels]
    x = np.arange(len(labels))
    
    plt.figure(figsize=(8, 5))
    plt.bar(x, maxs, color='skyblue', label='Max')
    plt.bar(x, mins, color='orange', label='Min')
    plt.xticks(x, labels, rotation=45)
    plt.title(title)
    plt.legend()
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()

plot_bar({k: data[k] for k in thermal}, "Thermal Properties", "thermal.jpeg")
plot_bar({k: data[k] for k in performance}, "Performance Indicators", "performance.jpeg")
plot_bar({k: data[k] for k in physical}, "Physical Characteristics", "physical.jpeg")

# Step 4: Radar chart
labels = list(data.keys())
values = [(data[k][0] + data[k][1]) / 2 for k in labels]
angles = [n / float(len(labels)) * 2 * pi for n in range(len(labels))]
values += values[:1]
angles += angles[:1]

plt.figure(figsize=(6, 6))
ax = plt.subplot(111, polar=True)
ax.plot(angles, values, color='green', linewidth=2)
ax.fill(angles, values, color='green', alpha=0.25)
plt.xticks(angles[:-1], labels)
plt.title("Radar Chart of Bioethanol Properties")
plt.tight_layout()
plt.savefig("radar.jpeg")
plt.close()

# Step 5: Create DOCX report
doc = Document()
doc.add_heading("Bioethanol Property Analysis", 0)

doc.add_heading("Thermal Properties", level=1)
doc.add_picture("thermal.jpeg", width=Inches(5.5))

doc.add_heading("Performance Indicators", level=1)
doc.add_picture("performance.jpeg", width=Inches(5.5))

doc.add_heading("Physical Characteristics", level=1)
doc.add_picture("physical.jpeg", width=Inches(5.5))

doc.add_heading("Radar Chart Overview", level=1)
doc.add_picture("radar.jpeg", width=Inches(5.5))

doc.add_paragraph("This report visualizes the key physical and chemical properties of bioethanol, highlighting its suitability as a fuel based on performance, thermal behavior, and physical traits.")

doc.save("bioethanol_analysis.docx")
