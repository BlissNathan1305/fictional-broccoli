import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import f_oneway, ttest_ind
from docx import Document
from docx.shared import Inches

# Step 1: Load the data
data = pd.DataFrame({
    'Time Point': ['72 hrs']*6 + ['4 days']*6 + ['5 days']*6 + ['6 days']*6,
    'Sample Type': ['Blank', 'Untreated', 'Acid treatment', 'Alkaline treatment', 'Enzyme', 'Acid w/o organism']*4,
    'Viable Cell Count': [
        0.000, 1.692, 1.747, 1.600, 1.557, 1.277,
        0.000, 1.752, 1.931, 1.810, 1.627, 1.422,
        0.000, 1.810, 1.985, 1.892, 1.714, 1.623,
        0.000, 1.943, 1.742, 1.648, 1.557, 1.610
    ],
    'pH': [
        7.0, 6.2, 5.5, 8.7, 6.4, 5.7,
        7.0, 5.8, 5.2, 7.6, 6.1, 5.1,
        7.0, 6.2, 5.5, 8.7, 6.4, 5.7,
        7.0, 5.9, 5.3, 8.1, 6.1, 5.5
    ],
    'Temperature': [30]*24
})

# Step 2: Descriptive statistics
desc_stats = data.groupby('Sample Type')['Viable Cell Count'].agg(['mean', 'std', 'min', 'max'])

# Step 3: ANOVA
groups = [data[data['Sample Type'] == sample]['Viable Cell Count'] for sample in data['Sample Type'].unique()]
anova_result = f_oneway(*groups)

# Step 4: Pairwise t-tests
sample_types = data['Sample Type'].unique()
ttest_results = []
for i in range(len(sample_types)):
    for j in range(i+1, len(sample_types)):
        group1 = data[data['Sample Type'] == sample_types[i]]['Viable Cell Count']
        group2 = data[data['Sample Type'] == sample_types[j]]['Viable Cell Count']
        t_stat, p_val = ttest_ind(group1, group2)
        ttest_results.append({
            'Group 1': sample_types[i],
            'Group 2': sample_types[j],
            't-statistic': round(t_stat, 3),
            'p-value': round(p_val, 5)
        })
ttest_df = pd.DataFrame(ttest_results)

# Step 5: Plotting
sns.set(style="whitegrid")

# Line plot
plt.figure(figsize=(8, 5))
sns.lineplot(data=data, x='Time Point', y='Viable Cell Count', hue='Sample Type', marker='o')
plt.title('Viable Cell Count Over Time')
plt.tight_layout()
plt.savefig('plantain_line.jpeg')
plt.close()

# Bar chart
plt.figure(figsize=(8, 5))
sns.barplot(data=data, x='Time Point', y='Viable Cell Count', hue='Sample Type')
plt.title('Bar Chart of Cell Counts')
plt.tight_layout()
plt.savefig('plantain_bar.jpeg')
plt.close()

# Box plot
plt.figure(figsize=(8, 5))
sns.boxplot(data=data, x='Sample Type', y='Viable Cell Count')
plt.title('Box Plot of Cell Counts by Treatment')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('plantain_box.jpeg')
plt.close()

# Step 6: Export to DOCX
doc = Document()
doc.add_heading('Statistical Analysis of Plantain Sap Fermentation', 0)

doc.add_heading('Descriptive Statistics', level=1)
for idx, row in desc_stats.iterrows():
    doc.add_paragraph(f"{idx}: Mean={row['mean']:.3f}, Std={row['std']:.3f}, Min={row['min']:.3f}, Max={row['max']:.3f}")

doc.add_heading('ANOVA Result', level=1)
doc.add_paragraph(f"F-statistic: {anova_result.statistic:.3f}, p-value: {anova_result.pvalue:.5f}")

doc.add_heading('Pairwise T-Test Results', level=1)
table = doc.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Group 1'
hdr_cells[1].text = 'Group 2'
hdr_cells[2].text = 't-statistic'
hdr_cells[3].text = 'p-value'
for _, row in ttest_df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Group 1']
    row_cells[1].text = row['Group 2']
    row_cells[2].text = str(row['t-statistic'])
    row_cells[3].text = str(row['p-value'])

doc.add_heading('Graphs', level=1)
doc.add_paragraph('Line Plot:')
doc.add_picture('plantain_line.jpeg', width=Inches(5))
doc.add_paragraph('Bar Chart:')
doc.add_picture('plantain_bar.jpeg', width=Inches(5))
doc.add_paragraph('Box Plot:')
doc.add_picture('plantain_box.jpeg', width=Inches(5))

doc.save('Plantain_Sap_Analysis.docx')
