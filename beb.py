# banana_sap_analysis.py
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# --- Input data ---
banana_sap_composition = {
    "Moisture (%)": 95.81,
    "Protein (%)": 1.75,
    "Fat/Lipid (%)": 0.23,
    "Fibre (%)": 0.00,
    "Ash (%)": 0.12,
    "Carbohydrate (%)": 2.08,
    "Energy (kcal/100g)": 17.39,
    "Lignin (%)": 0.01,
    "Hemicellulose (%)": 0.51,
    "Cellulose (%)": 0.58,
    "Sugar (%)": 5.13
}

# --- Setup output folder ---
output_dir = Path("banana_sap_analysis_outputs")
output_dir.mkdir(exist_ok=True)

# --- Basic dataframes ---
df_input = pd.DataFrame(
    list(banana_sap_composition.items()), columns=["Metric", "Value"]
).set_index("Metric")

# --- Derived metrics ---
moisture = banana_sap_composition["Moisture (%)"]
dry_matter = 100 - moisture

# Convert fresh basis to dry basis (except moisture & energy)
dry_basis = {}
for k, v in banana_sap_composition.items():
    if k not in ["Moisture (%)", "Energy (kcal/100g)"]:
        dry_basis[k + " (dry-basis %)"] = (v / dry_matter) * 100

# Energy per gram sugar
energy = banana_sap_composition["Energy (kcal/100g)"]
sugar = banana_sap_composition["Sugar (%)"]   # g/100 g fresh
energy_per_g_sugar = energy / sugar

# Ethanol yield (theoretical)
# 1 g sugar -> 0.511 g ethanol
ethanol_g_100g = sugar * 0.511
ethanol_ml_100g = ethanol_g_100g / 0.789  # density ethanol = 0.789 g/mL
ethanol_l_tonne = (ethanol_ml_100g * 10_000) / 1000  # L/tonne fresh

summary = {
    "Moisture (%)": moisture,
    "Dry matter (%)": dry_matter,
    "Sugar (g/100g fresh)": sugar,
    "Energy (kcal/100g fresh)": energy,
    "Energy per g sugar (kcal/g)": energy_per_g_sugar,
    "Ethanol (g/100g fresh)": ethanol_g_100g,
    "Ethanol (mL/100g fresh)": ethanol_ml_100g,
    "Ethanol (L/tonne fresh)": ethanol_l_tonne
}
df_summary = pd.DataFrame(list(summary.items()), columns=["Metric", "Value"]).set_index("Metric")

# --- PLOTS ---

# 1. Proximate composition
prox_keys = ["Moisture (%)", "Protein (%)", "Fat/Lipid (%)", "Fibre (%)",
             "Ash (%)", "Carbohydrate (%)", "Sugar (%)"]
prox_vals = [banana_sap_composition[k] for k in prox_keys]

plt.figure(figsize=(9,5))
plt.bar(prox_keys, prox_vals)
plt.title("Banana Sap — Proximate Composition")
plt.ylabel("Percent (%)")
plt.xticks(rotation=30, ha="right")
plt.tight_layout()
plt.savefig(output_dir/"proximate_composition.jpg", dpi=300)
plt.close()

# 2. Energy vs Sugar
plt.figure(figsize=(6,5))
plt.scatter([sugar], [energy], color="red")
plt.xlabel("Sugar (g/100 g fresh)")
plt.ylabel("Energy (kcal/100 g fresh)")
plt.title("Energy vs Sugar in Banana Sap")
plt.annotate(f"{energy_per_g_sugar:.2f} kcal/g sugar",
             (sugar, energy), xytext=(8,-12), textcoords="offset points")
plt.tight_layout()
plt.savefig(output_dir/"energy_vs_sugar.jpg", dpi=300)
plt.close()

# 3. Ethanol yield
eth_labels = ["g/100g", "mL/100g", "L/tonne"]
eth_values = [ethanol_g_100g, ethanol_ml_100g, ethanol_l_tonne]

plt.figure(figsize=(8,5))
plt.bar(eth_labels, eth_values, color="green")
plt.title("Theoretical Ethanol Yield from Banana Sap")
plt.ylabel("Amount")
plt.tight_layout()
plt.savefig(output_dir/"ethanol_yield.jpg", dpi=300)
plt.close()

# --- WORD REPORT ---
doc = Document()
doc.add_heading("Banana Sap Composition — Analysis & Bioethanol Metrics", level=1)

doc.add_heading("1. Methods", level=2)
doc.add_paragraph(
    "Proximate composition of banana sap was analyzed. Derived metrics include "
    "dry matter basis, energy-to-sugar ratio, and theoretical ethanol yield "
    "(assuming 0.511 g ethanol per g sugar, density 0.789 g/mL)."
)

doc.add_heading("2. Results (Tables)", level=2)
doc.add_paragraph("Input composition (fresh weight):")
t1 = doc.add_table(rows=1, cols=2)
t1.style = "Light List"
t1.rows[0].cells[0].text, t1.rows[0].cells[1].text = "Metric", "Value"
for m,v in banana_sap_composition.items():
    row = t1.add_row().cells
    row[0].text, row[1].text = m, str(v)

doc.add_paragraph("\nDerived summary:")
t2 = doc.add_table(rows=1, cols=2)
t2.style = "Light List"
t2.rows[0].cells[0].text, t2.rows[0].cells[1].text = "Metric", "Value"
for m,v in summary.items():
    row = t2.add_row().cells
    row[0].text = m
    row[1].text = f"{v:.4g}" if isinstance(v,float) else str(v)

doc.add_heading("3. Results (Figures)", level=2)
doc.add_picture(str(output_dir/"proximate_composition.jpg"), width=Inches(6))
doc.add_paragraph("Figure 1: Proximate composition of banana sap (fresh-weight %).")
doc.add_picture(str(output_dir/"energy_vs_sugar.jpg"), width=Inches(5))
doc.add_paragraph("Figure 2: Energy vs sugar content.")
doc.add_picture(str(output_dir/"ethanol_yield.jpg"), width=Inches(6))
doc.add_paragraph("Figure 3: Theoretical ethanol yield metrics.")

doc.add_heading("4. Discussion", level=2)
doc.add_paragraph(
    f"Banana sap contains {moisture:.2f}% moisture, leaving only {dry_matter:.2f}% "
    "dry matter. Sugar content is modest (5.13 g/100 g fresh), yielding "
    f"{ethanol_l_tonne:.2f} L ethanol per tonne (theoretical). "
    "This indicates that very large sap volumes are needed for significant ethanol production. "
    "The low energy density (%.2f kcal/g sugar) compared to other feedstocks reflects high dilution. "
    "Real-world ethanol yields would be lower due to process inefficiencies." % energy_per_g_sugar
)

doc.add_heading("5. Conclusions", level=2)
doc.add_paragraph(
    f"Banana sap has potential as a feedstock but its high water content and modest sugar "
    f"levels limit efficiency. Theoretical yield is about {ethanol_l_tonne:.2f} L ethanol per tonne."
)

doc.save(output_dir/"banana_sap_analysis_report.docx")

print("All files saved in:", output_dir.resolve())
