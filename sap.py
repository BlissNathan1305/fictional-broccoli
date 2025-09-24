import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

# === Step 1: Sample Data Setup ===

proximate_data = pd.DataFrame({
    'Component': ['Moisture', 'Fibre', 'Ash', 'Protein', 'Fat', 'Carbohydrate'],
    'Value (%)': [85.2, 3.1, 1.2, 1.5, 0.3, 8.7]
})

bioethanol_data = pd.DataFrame({
    'Component': ['Energy Content', 'Lignin', 'Cellulose', 'Hemicellulose', 'Reducing Sugars'],
    'Value': [16.5, 12.3, 25.4, 18.7, 9.8]
})

# === Step 2: Generate Graphs ===

# Proximate composition chart
plt.figure(figsize=(8, 5))
plt.bar(proximate_data['Component'], proximate_data['Value (%)'], color='skyblue')
plt.title('Proximate Composition of Banana Sap')
plt.ylabel('Percentage (%)')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('proximate_composition.jpeg')
plt.close()

# Bioethanol metrics chart
plt.figure(figsize=(8, 5))
plt.bar(bioethanol_data['Component'], bioethanol_data['Value'], color='orange')
plt.title('Bioethanol-Relevant Metrics from Banana Sap')
plt.ylabel('Value')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('bioethanol_metrics.jpeg')
plt.close()

# === Step 3: Create DOCX Report ===

doc = Document()

# Set default font to Times New Roman, 12 pt, black
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

doc.add_heading('Analysis of Banana Sap Composition for Bioethanol Production', 0)

# Proximate Composition Table
doc.add_heading('Proximate Composition', level=1)
table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value (%)'
for i, row in proximate_data.iterrows():
    cells = table1.add_row().cells
    cells[0].text = row['Component']
    cells[1].text = str(row['Value (%)'])

doc.add_picture('proximate_composition.jpeg', width=Inches(5.5))

# Bioethanol Metrics Table
doc.add_heading('Bioethanol-Relevant Metrics', level=1)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
for i, row in bioethanol_data.iterrows():
    cells = table2.add_row().cells
    cells[0].text = row['Component']
    cells[1].text = str(row['Value'])

doc.add_picture('bioethanol_metrics.jpeg', width=Inches(5.5))

# Discussion Section
doc.add_heading('Discussion', level=1)
doc.add_paragraph(
    "The proximate composition of banana sap reveals a high moisture content (85.2%), "
    "which may influence fermentation efficiency. The fibre and carbohydrate levels suggest "
    "potential for microbial activity, while low protein and fat indicate minimal nutritional interference.\n\n"
    "Bioethanol-relevant metrics show substantial cellulose (25.4%) and hemicellulose (18.7%) content, "
    "which are key substrates for ethanol production. Lignin (12.3%) may pose a challenge due to its resistance "
    "to enzymatic breakdown. The energy content (16.5 MJ/kg) supports its viability as a biofuel source.\n\n"
    "Overall, banana sap demonstrates promising characteristics for bioethanol production, though pretreatment "
    "strategies may be necessary to overcome lignin barriers."
)

doc.save('banana_sap_analysis.docx')
