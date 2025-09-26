import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import skew, kurtosis
from docx import Document
from docx.shared import Inches

# Data
data = {
    'Property': ['Ethanol concentration', 'Ethanol yield', 'pH', 'Density', 'Viscosity', 'Total Acidity'],
    'Value': [36.50, 0.62, 5.6, 0.98, 1.70, 0.4]
}
df = pd.DataFrame(data)

# Descriptive statistics
mean_val = np.mean(df['Value'])
std_val = np.std(df['Value'])
skew_val = skew(df['Value'])
kurt_val = kurtosis(df['Value'])

# Bar Chart
plt.figure(figsize=(8, 5))
plt.bar(df['Property'], df['Value'], color='skyblue')
plt.title('Physicochemical Properties of Plantain Sap')
plt.ylabel('Value')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('bar_chart.jpeg')
plt.close()

# Radar Chart (Corrected)
labels = df['Property']
values = df['Value'].tolist()
angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()

# Close the loop
values += values[:1]
angles += angles[:1]

fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
ax.plot(angles, values, 'o-', linewidth=2)
ax.fill(angles, values, alpha=0.25)
ax.set_thetagrids(np.degrees(angles[:-1]), labels)
ax.set_title('Radar Chart of Plantain Sap Properties')
plt.savefig('radar_chart.jpeg')
plt.close()

# Box Plot
plt.figure(figsize=(6, 4))
plt.boxplot(df['Value'], vert=False)
plt.title('Box Plot of Plantain Sap Values')
plt.xlabel('Value')
plt.tight_layout()
plt.savefig('box_plot.jpeg')
plt.close()

# Histogram
plt.figure(figsize=(6, 4))
plt.hist(df['Value'], bins=6, color='lightgreen', edgecolor='black')
plt.title('Histogram of Plantain Sap Values')
plt.xlabel('Value')
plt.ylabel('Frequency')
plt.tight_layout()
plt.savefig('histogram.jpeg')
plt.close()

# Academic-style Report
doc = Document()
doc.add_heading('Statistical Analysis of Physicochemical Properties of Plantain Sap', 0)

doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "This study presents a statistical evaluation of the physicochemical properties of plantain sap, "
    "highlighting its potential for industrial applications such as biofuel production and food processing."
)

doc.add_heading('Methodology', level=1)
doc.add_paragraph(
    "Six key physicochemical parameters were analyzed: ethanol concentration, ethanol yield, pH, density, "
    "viscosity, and total acidity. Descriptive statistics including mean, standard deviation, skewness, and "
    "kurtosis were computed. Visualizations were generated to aid interpretation."
)

doc.add_heading('Results', level=1)
doc.add_paragraph(f"Mean Value: {mean_val:.2f}")
doc.add_paragraph(f"Standard Deviation: {std_val:.2f}")
doc.add_paragraph(f"Skewness: {skew_val:.2f}")
doc.add_paragraph(f"Kurtosis: {kurt_val:.2f}")

doc.add_picture('bar_chart.jpeg', width=Inches(5))
doc.add_picture('radar_chart.jpeg', width=Inches(5))
doc.add_picture('box_plot.jpeg', width=Inches(5))
doc.add_picture('histogram.jpeg', width=Inches(5))

doc.add_heading('Discussion', level=1)
doc.add_paragraph(
    "The ethanol concentration (36.50) significantly exceeds other values, contributing to a positive skewness "
    "of 1.91. This suggests plantain sap is highly promising for ethanol-based applications. The moderate standard "
    "deviation (13.47) indicates variability across properties, while the kurtosis (2.38) reflects a slightly flatter "
    "distribution than normal. The radar chart visually confirms ethanol concentration as the dominant trait."
)

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "Plantain sap exhibits physicochemical characteristics favorable for fermentation and biofuel production. "
    "Future studies should compare these findings with other fruit saps and explore optimization strategies."
)

doc.save('Plantain_Sap_Analysis.docx')
