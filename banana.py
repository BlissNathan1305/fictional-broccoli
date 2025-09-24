import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import f_oneway, ttest_ind
from docx import Document
from docx.shared import Inches

# Step 1: Load the data
data = pd.DataFrame({
    'Time Point': ['72 hrs', '72 hrs', '72 hrs', '4 days', '4 days', '4 days',
                   '5 days', '5 days', '5 days', '6 days', '6 days', '6 days'],
    'Sample Type': ['Blank', 'Acid treatment', 'Alkaline treatment'] * 4,
    'Viable Cell Count': [0.000, 1.570, 1.796, 0.000, 1.610, 1.836,
                          0.000, 1.724, 1.912, 0.000, 1.411, 1.725],
    'pH': [7.0, 5.0, 8.0, 7.0, 4.8, 8.6, 7.0, 5.5, 8.2, 7.0, 5.0, 8.0],
    'Temperature': [30] * 12
})

# Step 2: Descriptive statistics
desc_stats = data.groupby('Sample Type')['Viable Cell Count'].agg(['mean', 'std', 'min', 'max'])

# Step 3: ANOVA
acid = data[data['Sample Type'] == 'Acid treatment']['Viable Cell Count']
alkaline = data[data['Sample Type'] == 'Alkaline treatment']['Viable Cell Count']
blank = data[data['Sample Type'] == 'Blank']['Viable Cell Count']
anova_result = f_oneway(acid, alkaline, blank)

# Step 4: T-tests
ttest_acid_alkaline = ttest_ind(acid, alkaline)
ttest_acid_blank = ttest_ind(acid, blank)
ttest_alkaline_blank = ttest_ind(alkaline, blank)

# Step 5: Plotting
sns.set(style="whitegrid")
plt.figure(figsize=(8, 5))
sns.lineplot(data=data, x='Time Point', y='Viable Cell Count', hue='Sample Type', marker='o')
plt.title('Viable Cell Count Over Time')
plt.savefig('line_plot.png')
plt.close()

plt.figure(figsize=(8, 5))
sns.barplot(data=data, x='Time Point', y='Viable Cell Count', hue='Sample Type')
plt.title('Bar Chart of Cell Counts')
plt.savefig('bar_chart.png')
plt.close()

plt.figure(figsize=(8, 5))
sns.boxplot(data=data, x='Sample Type', y='Viable Cell Count')
plt.title('Box Plot of Cell Counts by Treatment')
plt.savefig('box_plot.png')
plt.close()

# Step 6: Export to DOCX
doc = Document()
doc.add_heading('Statistical Analysis of Banana Sap Fermentation', 0)

doc.add_heading('Descriptive Statistics', level=1)
for idx, row in desc_stats.iterrows():
    doc.add_paragraph(f"{idx}: Mean={row['mean']:.3f}, Std={row['std']:.3f}, Min={row['min']:.3f}, Max={row['max']:.3f}")

doc.add_heading('ANOVA Result', level=1)
doc.add_paragraph(f"F-statistic: {anova_result.statistic:.3f}, p-value: {anova_result.pvalue:.5f}")

doc.add_heading('T-Test Results', level=1)
doc.add_paragraph(f"Acid vs Alkaline: t={ttest_acid_alkaline.statistic:.3f}, p={ttest_acid_alkaline.pvalue:.5f}")
doc.add_paragraph(f"Acid vs Blank: t={ttest_acid_blank.statistic:.3f}, p={ttest_acid_blank.pvalue:.5f}")
doc.add_paragraph(f"Alkaline vs Blank: t={ttest_alkaline_blank.statistic:.3f}, p={ttest_alkaline_blank.pvalue:.5f}")

doc.add_heading('Graphs', level=1)
doc.add_picture('line_plot.png', width=Inches(5))
doc.add_picture('bar_chart.png', width=Inches(5))
doc.add_picture('box_plot.png', width=Inches(5))

doc.save('Banana_Sap_Analysis.docx')
