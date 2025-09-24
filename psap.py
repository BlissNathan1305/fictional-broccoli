import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

# === Step 1: Plantain Sap Data Setup ===

proximate_data = pd.DataFrame({
    'Component': ['Moisture', 'Protein', 'Fat/Lipid', 'Fibre', 'Ash', 'Carbohydrate'],
    'Value (%)': [95.62, 1.63, 0.25, 0.0, 0.15, 2.34]
})

bioethanol_data = pd.DataFrame({
    'Component': ['Energy', 'Lignin', 'Hemicellulose', 'Cellulose', 'Sugar'],
    'Value': [18.13, 0.01, 0.52, 0.61, 5.13]
})

# === Step 2: Generate Graphs ===

# Proximate composition chart
plt.figure(figsize=(8, 5))
plt.bar(proximate_data['Component'], proximate_data['Value (%)'], color='mediumseagreen')
plt.title('Proximate Composition of Plantain Sap')
plt.ylabel('Percentage (%)')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('plantain_proximate_composition.jpeg')
plt.close()

# Bioethanol metrics chart
plt.figure(figsize=(8, 5))
plt.bar(bioethanol_data['Component'], bioethanol_data['Value'], color='coral')
plt.title('Bioethanol-Relevant Metrics from Plantain Sap')
plt.ylabel('Value')
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig('plantain_bioethanol_metrics.jpeg')
plt.close()

# === Step 3: Create DOCX Report ===

doc = Document()

# Set default font to Times New Roman, 12 pt, black
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

doc.add_heading('Analysis of Plantain Sap Composition for Bioethanol Production', 0)

# Proximate Composition Table
doc.add_heading('Proximate Composition', level=1)
table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Table Grid'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value (%)'
for i, row in proximate_data.iterrows():
    cells = table1.add_row().cells
    cells[0].text = row['Component']
    cells[1].text = str(row['Value (%)'])

doc.add_picture('plantain_proximate_composition.jpeg', width=Inches(5.5))

# Bioethanol Metrics Table
doc.add_heading('Bioethanol-Relevant Metrics', level=1)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Value'
for i, row in bioethanol_data.iterrows():
    cells = table2.add_row().cells
    cells[0].text = row['Component']
    cells[1].text = str(row['Value'])

doc.add_picture('plantain_bioethanol_metrics.jpeg', width=Inches(5.5))

# Discussion Section
doc.add_heading('Discussion', level=1)
doc.add_paragraph(
    "The proximate composition of plantain sap reveals an exceptionally high moisture content (95.62%), "
    "which may dilute fermentable substrates and affect fermentation efficiency. Protein (1.63%) and fat (0.25%) "
    "are present in low quantities, while fibre was not detected, suggesting minimal structural complexity.\n\n"
    "Carbohydrate content (2.34%) and sugar concentration (5.13 g/100g) indicate a modest potential for microbial fermentation. "
    "Bioethanol-relevant metrics such as cellulose (0.61%) and hemicellulose (0.52%) are relatively low, while lignin (0.01%) "
    "is nearly absent, which may reduce the need for pretreatment.\n\n"
    "The energy value (18.13 MJ/kg) supports its viability as a biofuel source. These findings align with observations by "
    "Rakhonde et al. (2019), who emphasized banana sap’s potential for ethanol recovery, and extend the scope to plantain sap, "
    "which shows similar but slightly more dilute characteristics.\n\n"
    "Overall, plantain sap presents a viable substrate for bioethanol production, particularly due to its low lignin content and moderate sugar levels, "
    "though its high moisture may require concentration or blending with other biomass sources."
)

# References
doc.add_heading('References', level=1)
doc.add_paragraph("1. Rakhonde MG, Waghmare GM, Garud HS. (2019). Production of bioethanol from banana scuitched sap. International Journal of Chemical Studies, 7(1): 2369–2371. https://www.chemijournal.com/archives/2019/vol7issue1/PartAO/7-1-556-203.pdf")
doc.add_paragraph("2. Wani S, Patil D. (2025). Nutritional and Biological Analysis of Nutrient-dense Banana Sap Water. International Journal of Environmental and Agriculture Research. https://ijoear.com/assets/articles_menuscripts/file/IJOEAR-JUL-2025-2.pdf")

doc.save('plantain_sap_analysis.docx')
