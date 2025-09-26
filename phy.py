import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import re

# ===============================
# Step 1: Create Dataset
# ===============================
data = {
    "Parameter": ["Ethanol concentration (%)", "Ethanol yield", "pH", 
                  "Density (g/cm³)", "Viscosity (mPa·s)", "Total acidity (%)"],
    "Plantain": [36.50, 0.62, 5.6, 0.98, 1.70, 0.40],
    "Banana": [32.70, 0.50, 5.4, 0.98, 1.30, 0.50]
}
df = pd.DataFrame(data)

# ===============================
# Helper function to clean filenames
# ===============================
def clean_filename(name, fruit):
    # Replace anything not alphanumeric with underscore
    safe_name = re.sub(r'[^A-Za-z0-9]+', '_', name)
    return f"{fruit}_{safe_name}.jpeg"

# ===============================
# Function to create report for one fruit
# ===============================
def generate_report(fruit):
    values = df[["Parameter", fruit]]
    
    # ===============================
    # Step 2: Generate Plots (saved as JPEG)
    # ===============================
    for i, row in values.iterrows():
        plt.figure()
        plt.bar([fruit], [row[fruit]], color="green" if fruit == "Plantain" else "yellow")
        plt.title(row["Parameter"])
        plt.ylabel("Value")
        fname = clean_filename(row["Parameter"], fruit)
        plt.savefig(fname, dpi=300)
        plt.close()
    
    # ===============================
    # Step 3: Create Word Document
    # ===============================
    doc = Document()
    doc.add_heading(f"Physicochemical Analysis of {fruit} Sap", level=1)
    
    doc.add_paragraph(
        f"This document presents a physicochemical analysis of {fruit.lower()} sap. "
        "Parameters measured include ethanol concentration, ethanol yield, pH, density, viscosity, "
        "and total acidity. Each result is visualized with plots and discussed in detail."
    )
    
    # Insert Table
    doc.add_heading("Table 1: Physicochemical Parameters", level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = fruit
    for i, row in values.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Parameter"])
        row_cells[1].text = str(row[fruit])
    
    # Insert Figures
    doc.add_heading("Figures", level=2)
    for i, row in values.iterrows():
        fig_path = clean_filename(row["Parameter"], fruit)
        doc.add_paragraph(row["Parameter"])
        doc.add_picture(fig_path, width=Inches(4.5))
    
    # ===============================
    # Step 4: Academic Discussion
    # ===============================
    if fruit == "Plantain":
        discussion = """
The physicochemical profile of plantain sap reveals a strong potential for fermentation 
and bioethanol production. The ethanol concentration (36.50%) is relatively high, 
suggesting that plantain sap possesses a rich pool of fermentable sugars. This is 
further supported by the ethanol yield of 0.62, which demonstrates favorable conversion 
efficiency. Such findings align with previous studies that report plantain cultivars 
as having higher starch and sugar reserves than bananas.

The measured pH of 5.6 indicates moderate acidity, which is within the tolerance range 
for common fermenting microorganisms such as Saccharomyces cerevisiae. This makes the 
sap conducive for controlled fermentation. The total acidity value (0.40%) confirms 
that plantain sap is less acidic compared to banana sap, potentially giving it an 
advantage in maintaining microbial stability without excessive sourness.

Viscosity is relatively high (1.70 mPa·s), which could be attributed to soluble fiber 
and polysaccharides. While this may slow fermentation kinetics, it also suggests 
greater nutritive density. Density (0.98 g/cm³) is comparable to most fruit juices, 
indicating balanced dissolved solids.

Overall, plantain sap demonstrates significant potential as a substrate for bioethanol 
production and as a raw material in beverage formulations. Its high ethanol yield and 
moderate acidity make it a promising candidate for industrial applications. Future 
investigations should include replicates, microbial profiling, and kinetic modeling 
to optimize fermentation processes.
"""
    else:  # Banana discussion
        discussion = """
The physicochemical characteristics of banana sap highlight its unique biochemical 
properties and potential industrial applications. The ethanol concentration (32.70%) 
is moderate and slightly lower than that of plantain, suggesting comparatively fewer 
fermentable sugars or less efficient fermentation dynamics. Ethanol yield (0.50) 
reinforces this, indicating reduced bioethanol conversion efficiency.

The sap’s pH of 5.4 falls within the acceptable fermentation range, yet it is slightly 
lower than plantain sap, reflecting higher acidity. This is further confirmed by the 
total acidity value of 0.50%, which surpasses plantain sap. Such acidity can influence 
the sensory properties of beverages, contributing to a sharper taste profile, while 
also impacting microbial activity during storage and fermentation.

Banana sap’s viscosity is relatively low (1.30 mPa·s), suggesting fewer soluble solids 
and a lighter texture compared to plantain. This property can facilitate faster 
fermentation kinetics, as lower viscosity typically allows better mass transfer of 
substrates and metabolites. The density value (0.98 g/cm³) is consistent with that 
of plantain sap, indicating similar overall soluble solid concentrations.

In summary, banana sap demonstrates moderate ethanol production potential but excels 
in acidity and lower viscosity. These attributes may make it more suitable for 
applications in food and beverage industries where acidity is desirable, rather than 
solely for bioethanol production. Further studies should evaluate sensory attributes, 
microbial ecology, and optimization strategies to enhance its industrial utility.
"""
    
    doc.add_heading("Discussion", level=2)
    doc.add_paragraph(discussion)
    
    # Save Document
    doc.save(f"{fruit}_Analysis.docx")

# ===============================
# Generate Reports for Both Fruits
# ===============================
generate_report("Plantain")
generate_report("Banana")
