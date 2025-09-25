import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from math import pi

# Step 1: Define banana sap data
banana_data = {
    "Octane number": (92, 105),
    "Flash point": (13, 15),
    "Density": (0.79, 0.80),
    "Viscosity": (1.1, 1.3),
    "Vapor pressure": (5.5, 6.0),
    "Calorific value": (26.5, 27.0),
    "Boiling point": (78, 78),
    "Freezing point": (-114, -95),
    "Autoignition temp.": (360, 370)
}

# Step 2: Categorize properties
thermal = ["Flash point", "Boiling point", "Freezing point", "Autoignition temp."]
performance = ["Octane number", "Calorific value"]
physical = ["Density", "Viscosity", "Vapor pressure"]

# Step 3: Plotting function
def plot_bar(data_subset, title, filename):
    labels = list(data_subset.keys())
    mins = [data_subset[k][0] for k in labels]
    maxs = [data_subset[k][1] for k in labels]
    x = np.arange(len(labels))
    
    plt.figure(figsize=(8, 5))
    plt.bar(x - 0.2, mins, width=0.4, color='orange', label='Min')
    plt.bar(x + 0.2, maxs, width=0.4, color='skyblue', label='Max')
    plt.xticks(x, labels, rotation=45)
    plt.title(title)
    plt.legend()
    plt.tight_layout()
    plt.savefig(filename)
    plt.close()

plot_bar({k: banana_data[k] for k in thermal}, "Thermal Properties", "banana_thermal.jpeg")
plot_bar({k: banana_data[k] for k in performance}, "Performance Indicators", "banana_performance.jpeg")
plot_bar({k: banana_data[k] for k in physical}, "Physical Characteristics", "banana_physical.jpeg")

# Step 4: Radar chart
labels = list(banana_data.keys())
values = [(banana_data[k][0] + banana_data[k][1]) / 2 for k in labels]
angles = [n / float(len(labels)) * 2 * pi for n in range(len(labels))]
values += values[:1]
angles += angles[:1]

plt.figure(figsize=(6, 6))
ax = plt.subplot(111, polar=True)
ax.plot(angles, values, color='green', linewidth=2)
ax.fill(angles, values, color='green', alpha=0.25)
plt.xticks(angles[:-1], labels)
plt.title("Radar Chart of Banana Sap Properties")
plt.tight_layout()
plt.savefig("banana_radar.jpeg")
plt.close()

# Step 5: Create DOCX report
doc = Document()
doc.add_heading("Banana Sap Property Analysis", 0)

doc.add_heading("Thermal Properties", level=1)
doc.add_picture("banana_thermal.jpeg", width=Inches(5.5))

doc.add_heading("Performance Indicators", level=1)
doc.add_picture("banana_performance.jpeg", width=Inches(5.5))

doc.add_heading("Physical Characteristics", level=1)
doc.add_picture("banana_physical.jpeg", width=Inches(5.5))

doc.add_heading("Radar Chart Overview", level=1)
doc.add_picture("banana_radar.jpeg", width=Inches(5.5))

doc.add_paragraph("This report visualizes the key physical and chemical properties of banana sap, highlighting its potential as a biofuel based on performance, thermal behavior, and physical traits.")

doc.save("banana_sap_analysis.docx")
