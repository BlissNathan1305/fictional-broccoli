import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Data
banana_sap_composition = {
    "Moisture (%)": 95.81,
    "Protein (%)": 1.75,
    "Fat/Lipid (%)": 0.23,
    "Fibre (%)": 0.00,
    "Ash (%)": 0.12,
    "Carbohydrate (%)": 2.08,
    "Energy (kcal/100g)": 17.39,
    "Lignin (%)": 0.01,
    "Hemicellulose (%)": 0.51,
    "Cellulose (%)": 0.58,
    "Sugar (%)": 5.13
}

# Plot 1: Proximate Composition
plt.figure(figsize=(10, 6))
plt.bar(banana_sap_composition.keys(), banana_sap_composition.values(), color='skyblue')
plt.xticks(rotation=45, ha='right')
plt.title("Proximate Composition of Banana Sap")
plt.ylabel("Amount")
plt.tight_layout()
plt.savefig("proximate_composition.jpeg")
plt.close()

# Plot 2: Bioethanol-Relevant Metrics
bioethanol_metrics = {k: banana_sap_composition[k] for k in ["Sugar (%)", "Cellulose (%)", "Hemicellulose (%)", "Lignin (%)"]}
plt.figure(figsize=(8, 5))
plt.bar(bioethanol_metrics.keys(), bioethanol_metrics.values(), color='lightgreen')
plt.title("Bioethanol-Relevant Metrics in Banana Sap")
plt.ylabel("Percentage (%)")
plt.tight_layout()
plt.savefig("bioethanol_metrics.jpeg")
plt.close()

# Create Word Document
doc = Document()
doc.add_heading("Banana Sap Composition Analysis", 0)

doc.add_paragraph("This document presents the proximate composition of banana sap and highlights key metrics relevant for bioethanol production.")

doc.add_heading("Proximate Composition", level=1)
doc.add_picture("proximate_composition.jpeg", width=Inches(6))

doc.add_heading("Bioethanol-Relevant Metrics", level=1)
doc.add_picture("bioethanol_metrics.jpeg", width=Inches(6))

doc.add_paragraph("Sugar, cellulose, hemicellulose, and lignin are important components for evaluating bioethanol potential. Banana sap shows promising sugar content with moderate cellulose and hemicellulose levels.")

doc.save("banana_sap_analysis.docx")
